# src/parsers.py

import pandas as pd
from typing import Optional, Dict, Any, IO
import json
import pytesseract
from PIL import Image
import openai
from docx import Document
from utils import retry_with_backoff

class AIFileParser:
    """
    Enhanced AI-powered file parser for various formats using OpenAI.
    Can handle text, CSV, Excel, DOCX, and image files.
    """

    def __init__(self, api_key: Optional[str] = None):
        """Initializes the OpenAI client if an API key is provided."""
        self.client = None
        if api_key:
            try:
                self.client = openai.OpenAI(api_key=api_key)
                self.model = "gpt-4o"
            except Exception as e:
                print(f"Failed to initialize OpenAI client for AIFileParser: {e}")

    def _get_file_preview(self, file: IO[bytes]) -> str:
        """Creates a text preview of a file (Excel, CSV, TXT) for the AI."""
        file.seek(0)
        try:  # Attempt to read as Excel
            df = pd.read_excel(file, header=None).head(15)
            return f"EXCEL FILE PREVIEW:\n{df.to_string()}"
        except Exception:
            file.seek(0)
            try:  # Fallback to reading as text (CSV/TXT)
                preview_content = file.read(2000)
                content = preview_content.decode('utf-8', errors='replace')
                return f"TEXT FILE PREVIEW:\n{content}"
            except Exception:
                return "Could not generate file preview."

    def _extract_text_from_image(self, file: IO[bytes]) -> str:
        """Extracts text from an image file using Tesseract OCR."""
        try:
            image = Image.open(file)
            return pytesseract.image_to_string(image)
        except Exception as e:
            return f"OCR Error: {e}"

    def _extract_text_from_docx(self, file: IO[bytes]) -> str:
        """Extracts full text from a DOCX file."""
        try:
            doc = Document(file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            return "\n".join(full_text)
        except Exception as e:
            return f"DOCX Error: {e}"

    @retry_with_backoff()
    def parse_rd_document(self, file: IO[bytes]) -> Dict[str, Any]:
        """
        Parses an R&D specification document (DOCX) to extract configuration data.
        Returns a dictionary with product info, risks, financials, and requirements.
        """
        if not self.client:
            return {"error": "AI client not initialized."}

        text_content = self._extract_text_from_docx(file)
        if not text_content or "Error" in text_content[:20]:
            return {"error": f"Could not extract text: {text_content}"}

        # Truncate if too long (approx 50k chars) to fit context window
        text_content = text_content[:50000]

        system_prompt = """
        You are a MedTech Product Manager. Analyze the provided R&D Product Specification document.
        Extract the following structured data to configure a Quality Management System:
        1. **Product Info**: Name, ALL SKUs associated with the product (as a list), Description (Intended Use/Features).
        2. **Financials**: Unit Cost, Sales Price (if found).
        3. **Risks**: Identify a list of risks for FMEA (Failure Mode, Effect, Cause).
        4. **Requirements**: Core User Needs and Technical Requirements.

        Return ONLY a valid JSON object with the following keys:
        {
            "product_name": "...",
            "skus": ["SKU1", "SKU2", ...],
            "description": "...",
            "unit_cost": 0.0,
            "sales_price": 0.0,
            "fmea_rows": [
                {"Potential Failure Mode": "...", "Potential Effect(s)": "...", "Potential Cause(s)": "..."}
            ],
            "user_needs": "...",
            "tech_requirements": "..."
        }
        """

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Here is the R&D Document Content:\n\n{text_content}"}
                ],
                response_format={"type": "json_object"}
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            return {"error": f"AI Parsing Error: {e}"}

    @retry_with_backoff()
    def analyze_file_structure(self, file: IO[bytes], target_sku: str) -> Dict[str, Any]:
        """
        Uses an AI model to analyze a file's structure and categorize its content.
        """
        if not self.client:
            return {"error": "AI client not initialized.", "filename": file.name}

        file_extension = file.name.split('.')[-1].lower()
        
        if file_extension == 'docx':
             # Special handling if a DOCX is uploaded in the generic uploader
             preview = self._extract_text_from_docx(file)[:2000]
        elif file_extension in ['png', 'jpg', 'jpeg']:
            ocr_text = self._extract_text_from_image(file)
            if not ocr_text:
                return {"error": "Could not extract text from image.", "filename": file.name}
            preview = f"OCR TEXT FROM IMAGE:\n{ocr_text[:2000]}"
        else:
            preview = self._get_file_preview(file)

        system_prompt = """
        You are a data analysis expert. Analyze the file preview to identify its content type.
        Return ONLY a single, valid JSON object.
        """
        user_prompt = f"""
        Analyze the following file preview for SKU: "{target_sku}".
        File Preview:
        {preview}

        Tasks:
        1. Identify content type: ['sales', 'returns', 'inspection', 'voice_of_customer', 'other'].
        2. Extract relevant data (total quantity, pass/fail, etc).
        3. Brief summary.

        Return JSON keys: "filename", "content_type", "key_data", "summary".
        """
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=1500,
                response_format={"type": "json_object"}
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            return {"error": f"AI Error: {e}", "filename": file.name}

    def extract_data(self, file: IO[bytes], analysis: Dict[str, Any], target_sku: str) -> Optional[pd.DataFrame]:
        """Extracts and standardizes data from a file."""
        content_type = analysis.get('content_type')
        if content_type not in ['sales', 'returns']:
            return None

        # Try to use the quantity directly extracted by the AI
        key_data = analysis.get('key_data', {})
        quantity = key_data.get('total_quantity', key_data.get('quantity'))

        if quantity is not None:
            try:
                clean_quantity = int(float(str(quantity).replace(',', '')))
                return pd.DataFrame([{'sku': target_sku, 'quantity': clean_quantity}])
            except (ValueError, TypeError):
                pass 

        # Fallback: Parse whole file
        file.seek(0)
        try:
            if file.name.endswith('.csv'):
                df = pd.read_csv(file)
            else:
                df = pd.read_excel(file)

            sku_col = next((col for col in df.columns if 'sku' in str(col).lower()), None)
            qty_col = next((col for col in df.columns if any(q in str(col).lower() for q in ['quantity', 'qty', 'sales', 'units', 'returned'])), None)

            if sku_col and qty_col:
                df.rename(columns={sku_col: 'sku', qty_col: 'quantity'}, inplace=True)
                sku_data = df[df['sku'] == target_sku]
                return sku_data[['sku', 'quantity']]

        except Exception as e:
            print(f"Fallback parsing failed for {file.name}: {e}")

        return None
