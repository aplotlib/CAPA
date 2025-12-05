# src/document_generator.py

from typing import Dict, Any, List, Optional
from io import BytesIO
from datetime import date
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentGenerator:
    """
    Generates comprehensive Word and Excel documents from session data.
    """

    def _add_main_table_row(self, table, heading: str, content: Any):
        """Helper to add a formatted row to a two-column table."""
        row_cells = table.add_row().cells
        p = row_cells[0].paragraphs[0]
        p.add_run(heading).bold = True
        # FIX: Handle None values gracefully
        row_cells[1].text = str(content) if content is not None else ''

    def _add_df_as_table(self, doc, df: pd.DataFrame, title: str):
        """Helper to add a pandas DataFrame as a formatted table in the document."""
        if df is None or df.empty:
            return
        doc.add_page_break()
        doc.add_heading(title, level=2)
        
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                # FIX: Handle None/NaN in dataframe cells
                row_cells[i].text = str(cell_value) if pd.notna(cell_value) else ""

    def _add_markdown_text(self, doc, text: str, title: str):
        """Adds text, potentially in Markdown, as a new section."""
        if not text:
            return
        doc.add_page_break()
        doc.add_heading(title, level=2)
        cleaned_text = text.replace('**', '').replace('`', '').replace('###', '').replace('##', '')
        doc.add_paragraph(cleaned_text)
    
    def _parse_and_add_markdown_table(self, doc, markdown_text: str, title: str):
        """Parses Markdown table and adds to doc."""
        if not markdown_text: return
        doc.add_page_break()
        doc.add_heading(title, level=2)
        lines = [line.strip() for line in markdown_text.strip().split('\n') if line.strip() and not line.strip().startswith(('###', '##'))]
        if not lines or '|' not in lines[0]:
            doc.add_paragraph(markdown_text) 
            return
        header_line = lines[0]
        headers = [h.strip() for h in header_line.strip('|').split('|')]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers): hdr_cells[i].text = header.replace('**', '')
        data_lines = [line for line in lines if '|' in line and '---' not in line and line != header_line]
        for line in data_lines:
            row_cells = table.add_row().cells
            cells = [c.strip().replace('**', '').replace('`', '') for c in line.strip('|').split('|')]
            for i, cell_text in enumerate(cells):
                if i < len(row_cells): row_cells[i].text = cell_text

    def generate_dashboard_docx(self, analysis_results: Dict[str, Any], product_info: Dict[str, str]) -> BytesIO:
        """
        Generates a report of the Dashboard / Mission Control metrics.
        """
        doc = Document()
        doc.add_heading("Dashboard Analytics Report", level=1)
        doc.add_paragraph(f"Date: {date.today().strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"Product: {product_info.get('name', 'N/A')} (SKU: {product_info.get('sku', 'N/A')})")
        
        if analysis_results:
            # Insights
            if 'insights' in analysis_results:
                self._add_markdown_text(doc, analysis_results['insights'], "AI Analysis Insights")
            
            # Summary Table
            summary_df = analysis_results.get('return_summary')
            if summary_df is not None and not summary_df.empty:
                self._add_df_as_table(doc, summary_df, "SKU Performance Summary")
                
            # Quality Metrics
            metrics = analysis_results.get('quality_metrics')
            if metrics:
                doc.add_heading("Quality Metrics", level=2)
                doc.add_paragraph(f"Risk Level: {metrics.get('risk_level', 'N/A')}")
                doc.add_paragraph(f"Quality Score: {metrics.get('quality_score', 'N/A')}/100")
        else:
            doc.add_paragraph("No analysis results available to export.")
            
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_summary_docx(self, session_data: Dict[str, Any], selected_sections: List[str]) -> BytesIO:
        """Generates a single comprehensive report based on selected sections."""
        doc = Document()
        doc.add_heading("ORION Project Summary Report", level=1)
        
        # Add Executive Summary if available
        if session_data.get('final_review_summary'):
             self._add_markdown_text(doc, session_data['final_review_summary'], "Executive Summary")
        
        if "CAPA Form" in selected_sections and session_data.get('capa_data'):
            self._generate_capa_form_section(doc, session_data['capa_data'])

        if "FMEA" in selected_sections and session_data.get('fmea_data') is not None:
            self._add_df_as_table(doc, session_data['fmea_data'], "Failure Mode and Effects Analysis (FMEA)")

        if "ISO 14971 Assessment" in selected_sections and session_data.get('risk_assessment'):
            self._parse_and_add_markdown_table(doc, session_data['risk_assessment'], "ISO 14971 Risk Assessment")
            
        if "URRA" in selected_sections:
            # Prefer DataFrame, fallback to markdown
            if isinstance(session_data.get('urra_df'), pd.DataFrame):
                self._add_df_as_table(doc, session_data['urra_df'], "Use-Related Risk Analysis (URRA)")
            elif session_data.get('urra'):
                self._parse_and_add_markdown_table(doc, session_data['urra'], "Use-Related Risk Analysis (URRA)")

        if "Vendor Email Draft" in selected_sections and session_data.get('vendor_email_draft'):
            self._add_markdown_text(doc, session_data['vendor_email_draft'], "Vendor Communication Draft")

        if "Human Factors Report" in selected_sections and session_data.get('human_factors_data'):
            self._generate_human_factors_section(doc, session_data['human_factors_data'])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _generate_capa_form_section(self, doc: Document, capa_data: Dict[str, Any]):
        """Generates the CAPA section."""
        doc.add_page_break()
        doc.add_heading("CAPA Report", level=1)
        
        table_info = doc.add_table(rows=0, cols=2)
        table_info.style = 'Table Grid'
        self._add_main_table_row(table_info, "CAPA Number:", capa_data.get('capa_number', ''))
        self._add_main_table_row(table_info, "Date:", str(capa_data.get('date', '')))
        self._add_main_table_row(table_info, "Prepared By:", capa_data.get('prepared_by', ''))
        self._add_main_table_row(table_info, "Product:", capa_data.get('product_name', ''))
        
        doc.add_paragraph() 
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.autofit = False 
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(5.0)

        fields = [
            ("Issue", "issue_description"),
            ("Immediate Actions/Corrections", "immediate_actions"),
            ("Root Cause", "root_cause"),
            ("Corrective Action", "corrective_action"),
            ("Implementation of Corrective Actions", "implementation_of_corrective_actions"),
            ("Preventive Action", "preventive_action"),
            ("Implementation of Preventive Actions", "implementation_of_preventive_actions"),
            ("Additional Comments", "additional_comments"),
            ("Effectiveness Check Plan", "effectiveness_verification_plan"),
            ("Effectiveness Check Findings", "effectiveness_check_findings"),
        ]

        for label, key in fields:
            self._add_main_table_row(table, label, capa_data.get(key, ""))

        doc.add_paragraph()
        doc.add_paragraph("__________________________\nSignature, Principal Investigator")
        doc.add_paragraph(f"Date: {capa_data.get('closure_date', '___________')}")

    def _generate_human_factors_section(self, doc: Document, hf_data: Dict[str, Any]):
        doc.add_page_break()
        doc.add_heading("Human Factors and Usability Engineering Report", level=2)
        section_map = {
            "Conclusion": "conclusion_statement", "Descriptions": "descriptions",
            "Device User Interface": "device_interface", "Known Use Problems": "known_problems",
            "Hazards and Risks Analysis": "hazards_analysis", "Preliminary Analyses": "preliminary_analyses",
            "Critical Tasks": "critical_tasks", "Validation Testing": "validation_testing"
        }
        for title, key in section_map.items():
            doc.add_heading(title, level=3)
            doc.add_paragraph(str(hf_data.get(key, "No data provided.")))

    def _add_section_heading(self, doc, text: str):
        p = doc.add_paragraph()
        p.add_run(text).bold = True

    def generate_project_charter_docx(self, charter_data: Dict[str, Any]) -> BytesIO:
        doc = Document()
        doc.add_heading(charter_data.get('project_name', 'Project Charter'), level=1)
        doc.add_paragraph(f"Date: {date.today().strftime('%Y-%m-%d')}")
        
        doc.add_heading("1. Project Overview & Business Case", level=2)
        self._add_section_heading(doc, "Problem Statement")
        doc.add_paragraph(charter_data.get('problem_statement', 'N/A'))
        self._add_section_heading(doc, "Project Goal")
        doc.add_paragraph(charter_data.get('project_goal', 'N/A'))
        self._add_section_heading(doc, "Project Scope")
        doc.add_paragraph(charter_data.get('scope', 'N/A'))

        doc.add_heading("2. Regulatory & Quality Strategy", level=2)
        self._add_section_heading(doc, "Device Classification (FDA)")
        doc.add_paragraph(charter_data.get('device_classification', 'N/A'))
        self._add_section_heading(doc, "Applicable Standards & Regulations")
        standards = charter_data.get('applicable_standards', [])
        if standards:
            for standard in standards:
                doc.add_paragraph(standard, style='List Bullet')
        else:
            doc.add_paragraph("N/A")

        doc.add_heading("3. Key Stakeholders", level=2)
        doc.add_paragraph(charter_data.get('stakeholders', 'N/A'))
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_capa_tracker_excel(self, session_data: Dict[str, Any]) -> BytesIO:
        capa_data = session_data.get('capa_data', {})
        status = "Closed" if capa_data.get('closure_date') else "Open"
        
        data = {
            'CAPA Number': [capa_data.get('capa_number', 'N/A')],
            'Status': [status],
            'Product SKU': [capa_data.get('product_name', 'N/A')],
            'Initiation Date': [capa_data.get('date')],
            'Closure Date': [capa_data.get('closure_date')],
            'Issue Description': [capa_data.get('issue_description', '')],
            'Root Cause': [capa_data.get('root_cause', '')],
            'Corrective Action': [capa_data.get('corrective_action', '')],
        }
        df = pd.DataFrame(data)

        for date_col in ['Initiation Date', 'Closure Date']:
            if not df[date_col].isna().all():
                df[date_col] = pd.to_datetime(df[date_col]).dt.strftime('%Y-%m-%d')

        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='CAPA_Tracker_Data')
            for i, col in enumerate(df.columns):
                col_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
                writer.sheets['CAPA_Tracker_Data'].set_column(i, i, col_len)
        output.seek(0)
        return output
        
    def generate_scar_docx(self, capa_data: Dict[str, Any], vendor_name: str) -> BytesIO:
        doc = Document()
        doc.add_heading(f"SCAR: {vendor_name}", level=1)
        doc.add_paragraph(f"Date: {date.today()}")
        doc.add_paragraph(f"Reference CAPA: {capa_data.get('capa_number', 'N/A')}")
        
        doc.add_heading("Issue Description", level=2)
        doc.add_paragraph(capa_data.get('issue_description', 'No description provided.'))
        
        doc.add_heading("Requirement", level=2)
        doc.add_paragraph("Please investigate the root cause of this issue and provide a corrective action plan within 14 days.")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
