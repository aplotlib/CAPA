"""
Medical Device CAPA Tool - Enhanced Version
A comprehensive quality management system for analyzing returns, sales data, 
and generating AI-powered CAPA documentation with multi-format support
"""

import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import json
from typing import Dict, List, Optional, Tuple
import base64
from io import BytesIO, StringIO
import openai
import anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from PIL import Image
import pytesseract
import logging
import requests
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import pdfplumber
import PyPDF2
import docx2txt
import openpyxl
import csv

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Page configuration
st.set_page_config(
    page_title="Medical Device CAPA Tool",
    page_icon="🏥",
    layout="wide"
)

# Professional medical device styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Professional medical theme */
    html, body, [class*="st-"] {
        font-family: 'Inter', sans-serif;
    }
    
    .main {
        background-color: #f8f9fa;
    }
    
    /* Header styling */
    .main-header {
        background: linear-gradient(135deg, #1e3a8a 0%, #2563eb 100%);
        color: white;
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .main-header h1 {
        margin: 0;
        font-size: 2.5rem;
        font-weight: 700;
    }
    
    .main-header p {
        margin: 0.5rem 0 0 0;
        font-size: 1.1rem;
        opacity: 0.9;
    }
    
    /* Metric cards */
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
        border-left: 4px solid;
        height: 100%;
    }
    
    .metric-card.critical {
        border-left-color: #dc2626;
    }
    
    .metric-card.warning {
        border-left-color: #f59e0b;
    }
    
    .metric-card.good {
        border-left-color: #10b981;
    }
    
    .metric-card.info {
        border-left-color: #3b82f6;
    }
    
    .metric-value {
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0;
    }
    
    .metric-label {
        font-size: 0.9rem;
        color: #6b7280;
        margin: 0;
    }
    
    .metric-change {
        font-size: 0.9rem;
        margin-top: 0.5rem;
    }
    
    /* Manual entry forms */
    .manual-entry-section {
        background: white;
        padding: 2rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
        margin-bottom: 1rem;
    }
    
    .manual-entry-header {
        font-size: 1.2rem;
        font-weight: 600;
        color: #1f2937;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #e5e7eb;
    }
    
    /* Status indicators */
    .status-badge {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 9999px;
        font-size: 0.875rem;
        font-weight: 500;
    }
    
    .status-critical {
        background-color: #fee2e2;
        color: #dc2626;
    }
    
    .status-warning {
        background-color: #fef3c7;
        color: #d97706;
    }
    
    .status-good {
        background-color: #d1fae5;
        color: #059669;
    }
    
    /* Section headers */
    .section-header {
        background: white;
        padding: 1rem 1.5rem;
        border-radius: 10px;
        margin: 2rem 0 1rem 0;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
        border-left: 4px solid #3b82f6;
    }
    
    .section-header h2 {
        margin: 0;
        font-size: 1.5rem;
        font-weight: 600;
        color: #1f2937;
    }
    
    /* Forms */
    .capa-form {
        background: white;
        padding: 2rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
    }
    
    .form-section {
        margin-bottom: 2rem;
        padding-bottom: 2rem;
        border-bottom: 1px solid #e5e7eb;
    }
    
    .form-section:last-child {
        border-bottom: none;
        margin-bottom: 0;
        padding-bottom: 0;
    }
    
    .form-label {
        font-weight: 600;
        color: #374151;
        margin-bottom: 0.5rem;
        display: block;
    }
    
    /* Alert boxes */
    .alert-box {
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        border: 1px solid;
    }
    
    .alert-critical {
        background-color: #fee2e2;
        border-color: #fecaca;
        color: #7f1d1d;
    }
    
    .alert-warning {
        background-color: #fef3c7;
        border-color: #fde68a;
        color: #78350f;
    }
    
    .alert-info {
        background-color: #dbeafe;
        border-color: #bfdbfe;
        color: #1e3a8a;
    }
    
    /* Screenshot gallery */
    .screenshot-gallery {
        display: grid;
        grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
        gap: 1rem;
        margin-top: 1rem;
    }
    
    .screenshot-item {
        background: white;
        padding: 0.5rem;
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    
    /* Buttons */
    .stButton > button {
        font-weight: 500;
        border-radius: 6px;
        padding: 0.5rem 1rem;
    }
    
    /* File uploader */
    .uploadedFile {
        border: 2px dashed #d1d5db;
        border-radius: 8px;
        padding: 1rem;
    }
    
    /* Google integration status */
    .google-status {
        padding: 0.5rem 1rem;
        border-radius: 8px;
        font-size: 0.875rem;
        margin: 0.5rem 0;
    }
    
    .google-connected {
        background-color: #d1fae5;
        color: #059669;
    }
    
    .google-disconnected {
        background-color: #fee2e2;
        color: #dc2626;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'sales_data' not in st.session_state:
    st.session_state.sales_data = None
if 'returns_data' not in st.session_state:
    st.session_state.returns_data = None
if 'metrics' not in st.session_state:
    st.session_state.metrics = {}
if 'capa_data' not in st.session_state:
    st.session_state.capa_data = {}
if 'screenshots' not in st.session_state:
    st.session_state.screenshots = []
if 'manual_entries' not in st.session_state:
    st.session_state.manual_entries = {'sales': [], 'returns': []}
if 'ai_client' not in st.session_state:
    st.session_state.ai_client = None
if 'google_creds' not in st.session_state:
    st.session_state.google_creds = None

# --- Enhanced File Processing ---
class UniversalFileProcessor:
    """Process various file formats including Google Docs/Sheets"""
    
    SUPPORTED_FORMATS = {
        'spreadsheet': ['.csv', '.tsv', '.xlsx', '.xls', '.ods'],
        'document': ['.pdf', '.docx', '.doc', '.txt', '.rtf'],
        'image': ['.png', '.jpg', '.jpeg', '.gif', '.bmp'],
        'google': ['application/vnd.google-apps.spreadsheet', 'application/vnd.google-apps.document']
    }
    
    @staticmethod
    def process_file(file, filename: str, file_type: str = None) -> pd.DataFrame:
        """Process any supported file type and return standardized DataFrame"""
        
        if not file_type:
            file_type = UniversalFileProcessor.detect_file_type(filename)
        
        try:
            # Spreadsheet formats
            if filename.endswith('.csv'):
                return pd.read_csv(file, encoding='utf-8', error_bad_lines=False)
            
            elif filename.endswith('.tsv') or filename.endswith('.txt'):
                # Try tab-delimited first
                try:
                    df = pd.read_csv(file, sep='\t', encoding='utf-8', error_bad_lines=False)
                    if len(df.columns) > 1:
                        return df
                except:
                    pass
                # Try comma-delimited
                return pd.read_csv(file, encoding='utf-8', error_bad_lines=False)
            
            elif filename.endswith(('.xlsx', '.xls')):
                # Read all sheets and combine
                excel_file = pd.ExcelFile(file)
                dfs = []
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    df['sheet_name'] = sheet_name
                    dfs.append(df)
                return pd.concat(dfs, ignore_index=True) if dfs else pd.DataFrame()
            
            # Document formats
            elif filename.endswith('.pdf'):
                return UniversalFileProcessor.extract_from_pdf(file)
            
            elif filename.endswith(('.docx', '.doc')):
                return UniversalFileProcessor.extract_from_word(file)
            
            # Image formats
            elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                return UniversalFileProcessor.extract_from_image(file)
            
            else:
                raise ValueError(f"Unsupported file format: {filename}")
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise
    
    @staticmethod
    def extract_from_pdf(pdf_file) -> pd.DataFrame:
        """Extract data from PDF using multiple methods"""
        all_data = []
        
        try:
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 1:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            all_data.append(df)
                    
                    # Extract text and parse
                    text = page.extract_text()
                    if text:
                        # Look for structured data patterns
                        lines = text.split('\n')
                        # Parse lines that look like data rows
                        data_rows = UniversalFileProcessor.parse_text_data(lines)
                        if data_rows:
                            all_data.append(pd.DataFrame(data_rows))
            
            if all_data:
                return pd.concat(all_data, ignore_index=True)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
        
        return pd.DataFrame()
    
    @staticmethod
    def extract_from_word(doc_file) -> pd.DataFrame:
        """Extract data from Word documents"""
        try:
            # Extract all text
            text = docx2txt.process(doc_file)
            
            # Look for tables in the document
            doc = Document(doc_file)
            all_tables = []
            
            for table in doc.tables:
                data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    data.append(row_data)
                
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    all_tables.append(df)
            
            if all_tables:
                return pd.concat(all_tables, ignore_index=True)
            
            # If no tables, try to parse structured text
            lines = text.split('\n')
            data_rows = UniversalFileProcessor.parse_text_data(lines)
            if data_rows:
                return pd.DataFrame(data_rows)
            
        except Exception as e:
            logger.error(f"Word extraction error: {e}")
        
        return pd.DataFrame()
    
    @staticmethod
    def extract_from_image(image_file) -> pd.DataFrame:
        """Extract text/data from images using OCR"""
        try:
            # Open image
            image = Image.open(image_file)
            
            # Extract text using OCR
            text = pytesseract.image_to_string(image)
            
            # Try to parse structured data
            lines = text.split('\n')
            data_rows = UniversalFileProcessor.parse_text_data(lines)
            
            if data_rows:
                return pd.DataFrame(data_rows)
            
            # If no structured data, return text as single column
            return pd.DataFrame({'extracted_text': [text]})
            
        except Exception as e:
            logger.error(f"Image extraction error: {e}")
            return pd.DataFrame()
    
    @staticmethod
    def parse_text_data(lines: List[str]) -> List[Dict]:
        """Parse text lines looking for structured data"""
        data_rows = []
        
        # Common patterns for data extraction
        patterns = {
            'order_id': re.compile(r'\b(\d{3}-\d{7}-\d{7})\b'),
            'sku': re.compile(r'\b([A-Z]{2,4}\d{3,6}[A-Z0-9]*)\b'),
            'quantity': re.compile(r'\b(\d+)\s*(?:units?|qty|quantity|pcs?)\b', re.I),
            'date': re.compile(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\w{3,9}\s+\d{1,2},?\s+\d{4})'),
            'price': re.compile(r'\$?\s*(\d+\.?\d*)')
        }
        
        for line in lines:
            if not line.strip():
                continue
            
            row_data = {}
            
            # Try to extract data using patterns
            for field, pattern in patterns.items():
                match = pattern.search(line)
                if match:
                    row_data[field] = match.group(1)
            
            if row_data:
                data_rows.append(row_data)
        
        return data_rows
    
    @staticmethod
    def detect_file_type(filename: str) -> str:
        """Detect file type from filename"""
        filename_lower = filename.lower()
        
        for category, extensions in UniversalFileProcessor.SUPPORTED_FORMATS.items():
            if any(filename_lower.endswith(ext) for ext in extensions):
                return category
        
        return 'unknown'

class GoogleIntegration:
    """Handle Google Docs and Sheets integration"""
    
    @staticmethod
    def authenticate():
        """Authenticate with Google services"""
        try:
            # Check for service account credentials in secrets
            if 'gcp_service_account' in st.secrets:
                creds = service_account.Credentials.from_service_account_info(
                    st.secrets['gcp_service_account'],
                    scopes=[
                        'https://www.googleapis.com/auth/drive.readonly',
                        'https://www.googleapis.com/auth/spreadsheets.readonly',
                        'https://www.googleapis.com/auth/documents.readonly'
                    ]
                )
                return creds
            return None
        except Exception as e:
            logger.error(f"Google authentication error: {e}")
            return None
    
    @staticmethod
    def get_google_sheet_data(sheet_id: str, range_name: str = 'A:Z') -> pd.DataFrame:
        """Fetch data from Google Sheets"""
        try:
            creds = GoogleIntegration.authenticate()
            if not creds:
                st.error("Google Sheets authentication failed")
                return pd.DataFrame()
            
            service = build('sheets', 'v4', credentials=creds)
            
            # Get data
            result = service.spreadsheets().values().get(
                spreadsheetId=sheet_id,
                range=range_name
            ).execute()
            
            values = result.get('values', [])
            
            if not values:
                return pd.DataFrame()
            
            # Convert to DataFrame
            df = pd.DataFrame(values[1:], columns=values[0])
            return df
            
        except Exception as e:
            logger.error(f"Error fetching Google Sheet: {e}")
            st.error(f"Failed to fetch Google Sheet: {str(e)}")
            return pd.DataFrame()
    
    @staticmethod
    def get_google_doc_data(doc_id: str) -> str:
        """Fetch content from Google Docs"""
        try:
            creds = GoogleIntegration.authenticate()
            if not creds:
                return ""
            
            service = build('docs', 'v1', credentials=creds)
            
            # Get document
            document = service.documents().get(documentId=doc_id).execute()
            
            # Extract text
            content = ''
            for element in document.get('body', {}).get('content', []):
                if 'paragraph' in element:
                    for text_element in element['paragraph'].get('elements', []):
                        if 'textRun' in text_element:
                            content += text_element['textRun'].get('content', '')
            
            return content
            
        except Exception as e:
            logger.error(f"Error fetching Google Doc: {e}")
            return ""

class ManualDataEntry:
    """Handle manual data entry forms"""
    
    @staticmethod
    def create_sales_entry_form():
        """Create form for manual sales data entry"""
        st.markdown('<div class="manual-entry-section">', unsafe_allow_html=True)
        st.markdown('<div class="manual-entry-header">📊 Manual Sales Entry</div>', unsafe_allow_html=True)
        
        with st.form("manual_sales_entry"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                sku = st.text_input("SKU", placeholder="e.g., LVA3100BLK")
                product_name = st.text_input("Product Name", placeholder="e.g., Wheelchair Bag")
                channel = st.selectbox("Sales Channel", ["Amazon", "B2B", "Direct", "Other"])
            
            with col2:
                quantity = st.number_input("Units Sold", min_value=0, value=0)
                date_from = st.date_input("Sales Period From", value=datetime.now() - timedelta(days=30))
                date_to = st.date_input("Sales Period To", value=datetime.now())
            
            with col3:
                order_count = st.number_input("Number of Orders", min_value=0, value=0)
                revenue = st.number_input("Total Revenue ($)", min_value=0.0, value=0.0, format="%.2f")
                notes = st.text_area("Notes", placeholder="Any additional information")
            
            submitted = st.form_submit_button("Add Sales Data", type="primary")
            
            if submitted and sku and quantity > 0:
                entry = {
                    'sku': sku,
                    'product_name': product_name,
                    'channel': channel,
                    'quantity': quantity,
                    'date_from': date_from,
                    'date_to': date_to,
                    'order_count': order_count,
                    'revenue': revenue,
                    'notes': notes,
                    'entry_date': datetime.now()
                }
                
                st.session_state.manual_entries['sales'].append(entry)
                st.success(f"✅ Added sales data for {sku}")
                
        st.markdown('</div>', unsafe_allow_html=True)
    
    @staticmethod
    def create_returns_entry_form():
        """Create form for manual returns data entry"""
        st.markdown('<div class="manual-entry-section">', unsafe_allow_html=True)
        st.markdown('<div class="manual-entry-header">📋 Manual Returns Entry</div>', unsafe_allow_html=True)
        
        with st.form("manual_returns_entry"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                sku = st.text_input("SKU", placeholder="e.g., LVA3100BLK", key="return_sku")
                order_id = st.text_input("Order ID", placeholder="e.g., 114-1234567-1234567")
                return_date = st.date_input("Return Date", value=datetime.now())
            
            with col2:
                quantity = st.number_input("Units Returned", min_value=1, value=1, key="return_qty")
                reason = st.selectbox("Return Reason", [
                    "Defective/Doesn't Work",
                    "Not Compatible",
                    "Wrong Item Sent",
                    "Not As Described",
                    "No Longer Needed",
                    "Bought By Mistake",
                    "Damaged in Shipping",
                    "Size/Fit Issues",
                    "Quality Issues",
                    "Other"
                ])
                
            with col3:
                customer_comments = st.text_area("Customer Comments", placeholder="Customer's return comments")
                internal_notes = st.text_area("Internal Notes", placeholder="Quality team notes")
            
            submitted = st.form_submit_button("Add Return Data", type="primary")
            
            if submitted and sku:
                entry = {
                    'sku': sku,
                    'order_id': order_id,
                    'return_date': return_date,
                    'quantity': quantity,
                    'reason': reason,
                    'customer_comments': customer_comments,
                    'internal_notes': internal_notes,
                    'entry_date': datetime.now()
                }
                
                st.session_state.manual_entries['returns'].append(entry)
                st.success(f"✅ Added return data for {sku}")
                
        st.markdown('</div>', unsafe_allow_html=True)
    
    @staticmethod
    def display_manual_entries():
        """Display current manual entries"""
        if st.session_state.manual_entries['sales'] or st.session_state.manual_entries['returns']:
            st.markdown('<div class="section-header"><h2>📝 Manual Entries Summary</h2></div>', unsafe_allow_html=True)
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.session_state.manual_entries['sales']:
                    st.subheader("Sales Entries")
                    sales_df = pd.DataFrame(st.session_state.manual_entries['sales'])
                    st.dataframe(sales_df[['sku', 'quantity', 'channel', 'date_from', 'date_to']], use_container_width=True)
            
            with col2:
                if st.session_state.manual_entries['returns']:
                    st.subheader("Returns Entries")
                    returns_df = pd.DataFrame(st.session_state.manual_entries['returns'])
                    st.dataframe(returns_df[['sku', 'quantity', 'reason', 'return_date']], use_container_width=True)

# --- Screenshot Processing ---
class ScreenshotProcessor:
    """Process and analyze screenshots"""
    
    @staticmethod
    def process_screenshots(screenshots: List) -> Dict:
        """Extract data from screenshots using OCR"""
        extracted_data = []
        
        for idx, screenshot in enumerate(screenshots):
            try:
                # Open image
                image = Image.open(screenshot)
                
                # Extract text
                text = pytesseract.image_to_string(image)
                
                # Try to identify data type
                data_type = ScreenshotProcessor.identify_screenshot_type(text)
                
                # Extract relevant information
                extracted_info = {
                    'screenshot_id': idx + 1,
                    'filename': screenshot.name,
                    'type': data_type,
                    'text': text,
                    'extracted_data': ScreenshotProcessor.extract_structured_data(text, data_type)
                }
                
                extracted_data.append(extracted_info)
                
            except Exception as e:
                logger.error(f"Error processing screenshot {screenshot.name}: {e}")
        
        return {
            'count': len(screenshots),
            'processed': len(extracted_data),
            'data': extracted_data
        }
    
    @staticmethod
    def identify_screenshot_type(text: str) -> str:
        """Identify the type of data in screenshot"""
        text_lower = text.lower()
        
        if 'return' in text_lower and ('reason' in text_lower or 'refund' in text_lower):
            return 'returns_data'
        elif 'order' in text_lower and 'quantity' in text_lower:
            return 'sales_data'
        elif 'defect' in text_lower or 'quality' in text_lower:
            return 'quality_report'
        elif 'complaint' in text_lower or 'issue' in text_lower:
            return 'complaint'
        else:
            return 'unknown'
    
    @staticmethod
    def extract_structured_data(text: str, data_type: str) -> Dict:
        """Extract structured data based on type"""
        extracted = {}
        
        # Common patterns
        patterns = {
            'order_id': re.compile(r'\b(\d{3}-\d{7}-\d{7})\b'),
            'sku': re.compile(r'\b([A-Z]{2,4}\d{3,6}[A-Z0-9]*)\b'),
            'date': re.compile(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'),
            'quantity': re.compile(r'(\d+)\s*(?:units?|qty|quantity)', re.I),
            'percentage': re.compile(r'(\d+\.?\d*)\s*%')
        }
        
        for field, pattern in patterns.items():
            matches = pattern.findall(text)
            if matches:
                extracted[field] = matches
        
        return extracted

# --- Helper Functions ---
class MetricsCalculator:
    """Calculate key quality metrics from sales and returns data"""
    
    @staticmethod
    def calculate_return_rate(sales_df, returns_df, group_by='sku'):
        """Calculate return rates by product"""
        try:
            # Aggregate sales by product
            sales_summary = sales_df.groupby(group_by).agg({
                'quantity': 'sum',
                'order_date': ['min', 'max', 'count']
            }).reset_index()
            sales_summary.columns = [group_by, 'total_sold', 'first_sale', 'last_sale', 'order_count']
            
            # Aggregate returns by product
            returns_summary = returns_df.groupby(group_by).agg({
                'quantity': 'sum',
                'return_date': 'count'
            }).reset_index()
            returns_summary.columns = [group_by, 'total_returned', 'return_count']
            
            # Merge and calculate rates
            summary = pd.merge(sales_summary, returns_summary, on=group_by, how='left')
            summary['total_returned'] = summary['total_returned'].fillna(0)
            summary['return_count'] = summary['return_count'].fillna(0)
            summary['return_rate'] = (summary['total_returned'] / summary['total_sold'] * 100).round(2)
            
            return summary
            
        except Exception as e:
            logger.error(f"Error calculating return rate: {e}")
            return pd.DataFrame()
    
    @staticmethod
    def calculate_sales_velocity(sales_df, periods=[30, 60, 90]):
        """Calculate sales velocity for different time periods"""
        try:
            velocities = {}
            current_date = pd.Timestamp.now()
            
            for period in periods:
                start_date = current_date - pd.Timedelta(days=period)
                period_sales = sales_df[sales_df['order_date'] >= start_date]
                
                velocity = period_sales.groupby('sku').agg({
                    'quantity': 'sum',
                    'order_id': 'count'
                }).reset_index()
                velocity.columns = ['sku', f'units_{period}d', f'orders_{period}d']
                velocity[f'daily_avg_{period}d'] = (velocity[f'units_{period}d'] / period).round(2)
                
                velocities[period] = velocity
            
            # Merge all periods
            result = velocities[periods[0]]
            for period in periods[1:]:
                result = pd.merge(result, velocities[period], on='sku', how='outer')
            
            return result.fillna(0)
            
        except Exception as e:
            logger.error(f"Error calculating sales velocity: {e}")
            return pd.DataFrame()
    
    @staticmethod
    def identify_quality_issues(returns_df, threshold=5):
        """Identify products with quality issues based on return patterns"""
        try:
            # Analyze return reasons
            reason_analysis = returns_df.groupby(['sku', 'reason']).agg({
                'quantity': 'sum',
                'return_id': 'count'
            }).reset_index()
            reason_analysis.columns = ['sku', 'reason', 'qty_returned', 'incident_count']
            
            # Identify top issues
            quality_indicators = ['defective', 'damaged', 'not_as_described', 'quality', 
                                'broken', 'malfunction', 'not_compatible']
            
            quality_issues = reason_analysis[
                reason_analysis['reason'].str.lower().str.contains('|'.join(quality_indicators), na=False)
            ]
            
            # Flag products exceeding threshold
            flagged_products = quality_issues[quality_issues['incident_count'] >= threshold]
            
            return flagged_products.sort_values('incident_count', ascending=False)
            
        except Exception as e:
            logger.error(f"Error identifying quality issues: {e}")
            return pd.DataFrame()

class DataProcessor:
    """Process various data formats from different channels"""
    
    @staticmethod
    def process_amazon_fba_returns(file_content):
        """Process Amazon FBA returns report"""
        try:
            # Read the tab-delimited file
            df = pd.read_csv(StringIO(file_content), sep='\t', parse_dates=['return-date'])
            
            # Standardize column names
            df = df.rename(columns={
                'return-date': 'return_date',
                'order-id': 'order_id',
                'sku': 'sku',
                'asin': 'asin',
                'product-name': 'product_name',
                'quantity': 'quantity',
                'reason': 'reason',
                'customer-comments': 'customer_comments',
                'detailed-disposition': 'disposition'
            })
            
            # Add return ID for tracking
            df['return_id'] = df.index + 1
            
            return df
            
        except Exception as e:
            logger.error(f"Error processing Amazon FBA returns: {e}")
            return None
    
    @staticmethod
    def process_sales_data(df):
        """Process sales data from various formats"""
        try:
            # Attempt to identify and standardize columns
            column_mapping = {
                'date': 'order_date',
                'order_date': 'order_date',
                'purchase_date': 'order_date',
                'sku': 'sku',
                'product_sku': 'sku',
                'item_sku': 'sku',
                'quantity': 'quantity',
                'qty': 'quantity',
                'units': 'quantity',
                'order_id': 'order_id',
                'order_number': 'order_id',
                'channel': 'channel',
                'marketplace': 'channel'
            }
            
            # Rename columns based on mapping
            for old_col, new_col in column_mapping.items():
                if old_col in df.columns:
                    df = df.rename(columns={old_col: new_col})
            
            # Ensure required columns exist
            required_cols = ['order_date', 'sku', 'quantity']
            missing_cols = [col for col in required_cols if col not in df.columns]
            if missing_cols:
                st.error(f"Missing required columns: {missing_cols}")
                return None
            
            # Convert date column
            df['order_date'] = pd.to_datetime(df['order_date'], errors='coerce')
            df['quantity'] = pd.to_numeric(df['quantity'], errors='coerce')
            
            # Remove invalid rows
            df = df.dropna(subset=['order_date', 'sku', 'quantity'])
            
            return df
            
        except Exception as e:
            logger.error(f"Error processing sales data: {e}")
            return None
    
    @staticmethod
    def combine_with_manual_entries(file_data: pd.DataFrame, manual_entries: List[Dict], 
                                  data_type: str) -> pd.DataFrame:
        """Combine file data with manual entries"""
        if not manual_entries:
            return file_data
        
        manual_df = pd.DataFrame(manual_entries)
        
        if data_type == 'sales':
            # Expand date ranges to individual records
            expanded_records = []
            for _, entry in manual_df.iterrows():
                # Create daily records for the period
                date_range = pd.date_range(entry['date_from'], entry['date_to'])
                daily_qty = entry['quantity'] / len(date_range)
                
                for date in date_range:
                    record = {
                        'order_date': date,
                        'sku': entry['sku'],
                        'quantity': daily_qty,
                        'channel': entry.get('channel', 'Manual'),
                        'order_id': f"MANUAL-{entry['sku']}-{date.strftime('%Y%m%d')}",
                        'source': 'manual_entry'
                    }
                    expanded_records.append(record)
            
            manual_processed = pd.DataFrame(expanded_records)
            
        else:  # returns
            manual_processed = manual_df.rename(columns={
                'return_date': 'return_date',
                'sku': 'sku',
                'quantity': 'quantity',
                'reason': 'reason',
                'customer_comments': 'customer_comments'
            })
            manual_processed['source'] = 'manual_entry'
            manual_processed['return_id'] = range(len(file_data) + 1, len(file_data) + len(manual_processed) + 1)
        
        # Combine data
        if file_data is not None and not file_data.empty:
            combined = pd.concat([file_data, manual_processed], ignore_index=True)
        else:
            combined = manual_processed
        
        return combined

class AIDocumentGenerator:
    """Generate CAPA documents using AI"""
    
    def __init__(self):
        self.openai_client = None
        self.anthropic_client = None
        self._initialize_clients()
    
    def _initialize_clients(self):
        """Initialize AI clients from Streamlit secrets"""
        try:
            if 'OPENAI_API_KEY' in st.secrets:
                openai.api_key = st.secrets['OPENAI_API_KEY']
                self.openai_client = openai
            
            if 'ANTHROPIC_API_KEY' in st.secrets:
                self.anthropic_client = anthropic.Anthropic(api_key=st.secrets['ANTHROPIC_API_KEY'])
                
        except Exception as e:
            logger.error(f"Error initializing AI clients: {e}")
    
    def generate_capa_document(self, capa_data, metrics, quality_issues, screenshots_data=None, use_anthropic=True):
        """Generate CAPA document using AI"""
        
        prompt = self._build_capa_prompt(capa_data, metrics, quality_issues, screenshots_data)
        
        try:
            if use_anthropic and self.anthropic_client:
                response = self.anthropic_client.messages.create(
                    model="claude-3-opus-20240229",
                    max_tokens=4000,
                    temperature=0,
                    messages=[{
                        "role": "user",
                        "content": prompt
                    }]
                )
                return response.content[0].text
                
            elif self.openai_client:
                response = openai.ChatCompletion.create(
                    model="gpt-4",
                    messages=[{
                        "role": "system",
                        "content": "You are a medical device quality management expert creating formal CAPA documentation."
                    }, {
                        "role": "user",
                        "content": prompt
                    }],
                    temperature=0,
                    max_tokens=4000
                )
                return response.choices[0].message.content
                
            else:
                return None
                
        except Exception as e:
            logger.error(f"Error generating CAPA document: {e}")
            return None
    
    def _build_capa_prompt(self, capa_data, metrics, quality_issues, screenshots_data):
        """Build prompt for CAPA generation"""
        
        # Include screenshot insights if available
        screenshot_insights = ""
        if screenshots_data and screenshots_data.get('data'):
            screenshot_insights = "\n\nSCREENSHOT EVIDENCE:"
            for item in screenshots_data['data']:
                if item.get('extracted_data'):
                    screenshot_insights += f"\n- {item['type']}: {item.get('extracted_data', {})}"
        
        prompt = f"""Generate a formal CAPA (Corrective and Preventive Action) report for a medical device quality issue.

CAPA DATA:
- CAPA Number: {capa_data.get('capa_number', 'TBD')}
- Date: {datetime.now().strftime('%Y-%m-%d')}
- Product: {capa_data.get('product', 'TBD')}
- SKU: {capa_data.get('sku', 'TBD')}

ISSUE DESCRIPTION:
{capa_data.get('issue_description', 'TBD')}

METRICS SUMMARY:
- Overall Return Rate: {metrics.get('overall_return_rate', 0):.2f}%
- Units Returned: {metrics.get('total_returns', 0)}
- Units Sold: {metrics.get('total_sales', 0)}
- Time Period: {metrics.get('period', 'Last 90 days')}

TOP QUALITY ISSUES:
{quality_issues.to_string() if not quality_issues.empty else 'No specific quality issues identified'}

{screenshot_insights}

PROPOSED CORRECTIVE ACTION:
{capa_data.get('corrective_action', 'TBD')}

Please generate a comprehensive CAPA report following medical device quality standards that includes:

1. ISSUE IDENTIFICATION
   - Clear problem statement
   - Impact assessment
   - Risk classification

2. ROOT CAUSE ANALYSIS
   - Investigation methodology
   - Root cause findings
   - Contributing factors

3. CORRECTIVE ACTIONS
   - Immediate actions taken
   - Long-term corrective measures
   - Implementation timeline
   - Responsible parties

4. PREVENTIVE ACTIONS
   - Systemic improvements
   - Process changes
   - Training requirements
   - Monitoring plans

5. EFFECTIVENESS VERIFICATION
   - Success criteria
   - Verification methods
   - Timeline for effectiveness checks
   - Key performance indicators

6. RISK ASSESSMENT
   - Patient safety impact
   - Regulatory considerations
   - Business impact

Format the response as a professional CAPA document suitable for regulatory submission."""
        
        return prompt
    
    def export_to_docx(self, content, capa_data, include_screenshots=False, screenshots=[]):
        """Export CAPA content to Word document"""
        try:
            doc = Document()
            
            # Add header
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = "CORRECTIVE AND PREVENTIVE ACTION (CAPA) REPORT"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add title
            title = doc.add_heading('CAPA Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add CAPA details
            doc.add_paragraph(f"CAPA Number: {capa_data.get('capa_number', 'TBD')}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph(f"Product: {capa_data.get('product', 'TBD')}")
            doc.add_paragraph(f"Prepared By: {capa_data.get('prepared_by', 'Quality Team')}")
            
            doc.add_paragraph()
            
            # Add content sections
            sections = content.split('\n\n')
            for section in sections:
                if section.strip():
                    if section.strip().isupper() or section.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.')):
                        # Section header
                        doc.add_heading(section.strip(), level=1)
                    else:
                        # Regular paragraph
                        doc.add_paragraph(section.strip())
            
            # Add screenshots if requested
            if include_screenshots and screenshots:
                doc.add_page_break()
                doc.add_heading('Supporting Evidence - Screenshots', level=1)
                
                for idx, screenshot in enumerate(screenshots):
                    try:
                        # Add screenshot
                        doc.add_paragraph(f"Screenshot {idx + 1}: {screenshot.name}")
                        
                        # Convert to image and add to document
                        image = Image.open(screenshot)
                        
                        # Save to BytesIO
                        img_buffer = BytesIO()
                        image.save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        
                        # Add to document (max width 6 inches)
                        doc.add_picture(img_buffer, width=Inches(6))
                        doc.add_paragraph()
                        
                    except Exception as e:
                        logger.error(f"Error adding screenshot {idx + 1}: {e}")
                        doc.add_paragraph(f"[Error loading screenshot {idx + 1}]")
            
            # Add signature section
            doc.add_page_break()
            doc.add_heading('Approvals', level=1)
            
            table = doc.add_table(rows=4, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            cells = table.rows[0].cells
            cells[0].text = 'Role'
            cells[1].text = 'Name / Signature'
            cells[2].text = 'Date'
            
            # Signature rows
            roles = ['Prepared By', 'Reviewed By', 'Approved By']
            for i, role in enumerate(roles, 1):
                cells = table.rows[i].cells
                cells[0].text = role
                cells[1].text = '_' * 30
                cells[2].text = '_' * 15
            
            # Save to BytesIO
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            
            return bio
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            return None

# --- UI Components ---
def display_header():
    """Display application header"""
    st.markdown("""
    <div class="main-header">
        <h1>🏥 Medical Device CAPA Tool</h1>
        <p>Quality Management System for Returns Analysis and CAPA Generation</p>
    </div>
    """, unsafe_allow_html=True)

def display_metrics_dashboard(metrics):
    """Display key metrics dashboard"""
    col1, col2, col3, col4 = st.columns(4)
    
    # Return Rate
    with col1:
        return_rate = metrics.get('overall_return_rate', 0)
        if return_rate > 10:
            status_class = "critical"
            status_color = "#dc2626"
        elif return_rate > 5:
            status_class = "warning"
            status_color = "#f59e0b"
        else:
            status_class = "good"
            status_color = "#10b981"
        
        st.markdown(f"""
        <div class="metric-card {status_class}">
            <p class="metric-label">Overall Return Rate</p>
            <p class="metric-value" style="color: {status_color}">{return_rate:.1f}%</p>
            <p class="metric-change">Target: <5%</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Total Returns
    with col2:
        total_returns = metrics.get('total_returns', 0)
        st.markdown(f"""
        <div class="metric-card info">
            <p class="metric-label">Total Returns</p>
            <p class="metric-value" style="color: #3b82f6">{total_returns:,}</p>
            <p class="metric-change">{metrics.get('period', 'All Time')}</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Products Affected
    with col3:
        products_affected = metrics.get('products_affected', 0)
        st.markdown(f"""
        <div class="metric-card info">
            <p class="metric-label">Products Affected</p>
            <p class="metric-value" style="color: #3b82f6">{products_affected}</p>
            <p class="metric-change">Unique SKUs</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Quality Issues
    with col4:
        quality_issues = metrics.get('quality_issues', 0)
        if quality_issues > 0:
            status_class = "warning"
            status_color = "#f59e0b"
        else:
            status_class = "good"
            status_color = "#10b981"
        
        st.markdown(f"""
        <div class="metric-card {status_class}">
            <p class="metric-label">Quality Issues</p>
            <p class="metric-value" style="color: {status_color}">{quality_issues}</p>
            <p class="metric-change">Flagged Products</p>
        </div>
        """, unsafe_allow_html=True)

def display_quality_alerts(quality_issues):
    """Display quality alerts for products with issues"""
    if quality_issues.empty:
        return
    
    st.markdown('<div class="section-header"><h2>⚠️ Quality Alerts</h2></div>', unsafe_allow_html=True)
    
    for _, issue in quality_issues.iterrows():
        severity = "critical" if issue['incident_count'] >= 10 else "warning"
        st.markdown(f"""
        <div class="alert-box alert-{severity}">
            <strong>SKU: {issue['sku']}</strong><br>
            Reason: {issue['reason']}<br>
            Incidents: {issue['incident_count']} | Units: {issue['qty_returned']}
        </div>
        """, unsafe_allow_html=True)

def display_screenshot_gallery(screenshots):
    """Display uploaded screenshots in a gallery"""
    if not screenshots:
        return
    
    st.markdown('<div class="section-header"><h2>📸 Uploaded Screenshots</h2></div>', unsafe_allow_html=True)
    
    # Process screenshots
    screenshot_data = ScreenshotProcessor.process_screenshots(screenshots)
    st.session_state.screenshot_data = screenshot_data
    
    # Display gallery
    cols = st.columns(4)
    for idx, (screenshot, info) in enumerate(zip(screenshots, screenshot_data['data'])):
        with cols[idx % 4]:
            st.image(screenshot, caption=f"{info['filename']} - {info['type']}", use_container_width=True)
            
            # Show extracted data if any
            if info.get('extracted_data'):
                with st.expander("Extracted Data"):
                    st.json(info['extracted_data'])

# --- Main Application ---
def main():
    display_header()
    
    # Sidebar for data input options
    with st.sidebar:
        st.header("📁 Data Input Options")
        
        input_method = st.radio(
            "Choose input method:",
            ["File Upload", "Manual Entry", "Google Integration", "Combined"]
        )
        
        if input_method in ["File Upload", "Combined"]:
            st.subheader("📤 File Upload")
            
            st.markdown("**Sales Data**")
            sales_files = st.file_uploader(
                "Upload sales data",
                type=['csv', 'tsv', 'txt', 'xlsx', 'xls', 'pdf', 'docx'],
                accept_multiple_files=True,
                help="Supports CSV, TSV, TXT, Excel, PDF, Word formats"
            )
            
            st.markdown("**Returns Data**")
            returns_files = st.file_uploader(
                "Upload returns data",
                type=['csv', 'tsv', 'txt', 'xlsx', 'xls', 'pdf', 'docx'],
                accept_multiple_files=True,
                help="Amazon FBA returns, PDFs, or any supported format"
            )
            
            st.markdown("**Screenshots**")
            screenshots = st.file_uploader(
                "Upload screenshots",
                type=['png', 'jpg', 'jpeg', 'gif', 'bmp'],
                accept_multiple_files=True,
                help="Upload any relevant screenshots or images"
            )
            
            if st.button("🔄 Process Files", type="primary"):
                process_uploaded_files(sales_files, returns_files, screenshots)
        
        if input_method in ["Google Integration", "Combined"]:
            st.subheader("🔗 Google Integration")
            
            # Check authentication status
            if st.session_state.google_creds:
                st.markdown('<div class="google-status google-connected">✅ Google Connected</div>', unsafe_allow_html=True)
            else:
                st.markdown('<div class="google-status google-disconnected">❌ Not Connected</div>', unsafe_allow_html=True)
            
            # Google Sheets
            sheet_url = st.text_input("Google Sheets URL", placeholder="https://docs.google.com/spreadsheets/d/...")
            if sheet_url and st.button("Import Sheet"):
                import_google_sheet(sheet_url)
            
            # Google Docs
            doc_url = st.text_input("Google Docs URL", placeholder="https://docs.google.com/document/d/...")
            if doc_url and st.button("Import Doc"):
                import_google_doc(doc_url)
    
    # Main content area
    if input_method in ["Manual Entry", "Combined"]:
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["📝 Manual Entry", "📊 Dashboard", "📈 Analysis", "📋 CAPA Form", "📄 Documents"])
    else:
        tab1, tab2, tab3, tab4 = st.tabs(["📊 Dashboard", "📈 Analysis", "📋 CAPA Form", "📄 Documents"])
    
    if input_method in ["Manual Entry", "Combined"]:
        with tab1:
            st.header("Manual Data Entry")
            
            # Sales entry
            ManualDataEntry.create_sales_entry_form()
            
            # Returns entry
            ManualDataEntry.create_returns_entry_form()
            
            # Display current entries
            ManualDataEntry.display_manual_entries()
            
            # Process manual data button
            if st.button("📊 Calculate Metrics from Manual Data", type="primary"):
                process_manual_data()
    
    # Dashboard tab
    dashboard_tab = tab2 if input_method in ["Manual Entry", "Combined"] else tab1
    with dashboard_tab:
        if st.session_state.metrics:
            display_metrics_dashboard(st.session_state.metrics)
            
            # Return rate trend
            if st.session_state.returns_data is not None:
                st.markdown('<div class="section-header"><h2>📈 Return Rate Trend</h2></div>', unsafe_allow_html=True)
                display_return_trend(st.session_state.returns_data)
            
            # Quality alerts
            if 'quality_issues' in st.session_state and not st.session_state.quality_issues.empty:
                display_quality_alerts(st.session_state.quality_issues)
            
            # Screenshots gallery
            if st.session_state.screenshots:
                display_screenshot_gallery(st.session_state.screenshots)
        else:
            st.info("📤 Please upload data or enter manual data to view metrics")
    
    # Analysis tab
    analysis_tab = tab3 if input_method in ["Manual Entry", "Combined"] else tab2
    with analysis_tab:
        if st.session_state.sales_data is not None and st.session_state.returns_data is not None:
            display_detailed_analysis()
        else:
            st.info("📤 Please provide data to view analysis")
    
    # CAPA Form tab
    capa_tab = tab4 if input_method in ["Manual Entry", "Combined"] else tab3
    with capa_tab:
        display_capa_form()
    
    # Documents tab
    docs_tab = tab5 if input_method in ["Manual Entry", "Combined"] else tab4
    with docs_tab:
        display_document_generation()

def process_uploaded_files(sales_files, returns_files, screenshots):
    """Process all uploaded files"""
    
    with st.spinner("Processing files..."):
        processor = UniversalFileProcessor()
        
        # Process sales files
        all_sales_data = []
        if sales_files:
            for file in sales_files:
                try:
                    df = processor.process_file(file, file.name)
                    if not df.empty:
                        sales_df = DataProcessor.process_sales_data(df)
                        if sales_df is not None:
                            all_sales_data.append(sales_df)
                            st.success(f"✅ Processed {file.name}: {len(sales_df)} sales records")
                except Exception as e:
                    st.error(f"Error processing {file.name}: {e}")
            
            if all_sales_data:
                combined_sales = pd.concat(all_sales_data, ignore_index=True)
                # Combine with manual entries
                combined_sales = DataProcessor.combine_with_manual_entries(
                    combined_sales, 
                    st.session_state.manual_entries['sales'],
                    'sales'
                )
                st.session_state.sales_data = combined_sales
                st.success(f"✅ Total sales records: {len(combined_sales)}")
        
        # Process returns files
        all_returns_data = []
        if returns_files:
            for file in returns_files:
                try:
                    # Check if it's an FBA returns file
                    if file.name.endswith('.txt'):
                        content = file.read().decode('utf-8')
                        if 'return-date' in content and 'order-id' in content:
                            df = DataProcessor.process_amazon_fba_returns(content)
                        else:
                            file.seek(0)
                            df = processor.process_file(file, file.name)
                    else:
                        df = processor.process_file(file, file.name)
                    
                    if df is not None and not df.empty:
                        all_returns_data.append(df)
                        st.success(f"✅ Processed {file.name}: {len(df)} return records")
                except Exception as e:
                    st.error(f"Error processing {file.name}: {e}")
            
            if all_returns_data:
                combined_returns = pd.concat(all_returns_data, ignore_index=True)
                # Combine with manual entries
                combined_returns = DataProcessor.combine_with_manual_entries(
                    combined_returns,
                    st.session_state.manual_entries['returns'],
                    'returns'
                )
                st.session_state.returns_data = combined_returns
                st.success(f"✅ Total return records: {len(combined_returns)}")
        
        # Store screenshots
        if screenshots:
            st.session_state.screenshots = screenshots
            st.success(f"✅ Loaded {len(screenshots)} screenshots")
        
        # Calculate metrics
        if st.session_state.sales_data is not None and st.session_state.returns_data is not None:
            calculate_metrics()

def process_manual_data():
    """Process manual entries and calculate metrics"""
    
    # Convert manual entries to DataFrames
    if st.session_state.manual_entries['sales']:
        sales_df = DataProcessor.combine_with_manual_entries(
            None,
            st.session_state.manual_entries['sales'],
            'sales'
        )
        st.session_state.sales_data = sales_df
    
    if st.session_state.manual_entries['returns']:
        returns_df = DataProcessor.combine_with_manual_entries(
            None,
            st.session_state.manual_entries['returns'],
            'returns'
        )
        st.session_state.returns_data = returns_df
    
    # Calculate metrics
    if st.session_state.sales_data is not None and st.session_state.returns_data is not None:
        calculate_metrics()
        st.success("✅ Metrics calculated from manual data")

def import_google_sheet(sheet_url):
    """Import data from Google Sheets"""
    try:
        # Extract sheet ID from URL
        import re
        match = re.search(r'/spreadsheets/d/([a-zA-Z0-9-_]+)', sheet_url)
        if not match:
            st.error("Invalid Google Sheets URL")
            return
        
        sheet_id = match.group(1)
        
        # Fetch data
        df = GoogleIntegration.get_google_sheet_data(sheet_id)
        
        if not df.empty:
            # Try to determine if it's sales or returns data
            if any(col in df.columns for col in ['return_date', 'return-date', 'reason']):
                st.session_state.returns_data = DataProcessor.process_sales_data(df)
                st.success(f"✅ Imported {len(df)} return records from Google Sheets")
            else:
                st.session_state.sales_data = DataProcessor.process_sales_data(df)
                st.success(f"✅ Imported {len(df)} sales records from Google Sheets")
            
            calculate_metrics()
        
    except Exception as e:
        st.error(f"Failed to import Google Sheet: {e}")

def import_google_doc(doc_url):
    """Import data from Google Docs"""
    try:
        # Extract doc ID from URL
        import re
        match = re.search(r'/document/d/([a-zA-Z0-9-_]+)', doc_url)
        if not match:
            st.error("Invalid Google Docs URL")
            return
        
        doc_id = match.group(1)
        
        # Fetch content
        content = GoogleIntegration.get_google_doc_data(doc_id)
        
        if content:
            st.text_area("Document Content", content, height=300)
            st.info("Document content extracted. You may need to manually copy relevant data.")
        
    except Exception as e:
        st.error(f"Failed to import Google Doc: {e}")

def calculate_metrics():
    """Calculate all metrics from loaded data"""
    
    calc = MetricsCalculator()
    
    # Calculate return rates
    return_summary = calc.calculate_return_rate(
        st.session_state.sales_data, 
        st.session_state.returns_data
    )
    
    # Calculate sales velocity
    velocity = calc.calculate_sales_velocity(st.session_state.sales_data)
    
    # Identify quality issues
    quality_issues = calc.identify_quality_issues(st.session_state.returns_data)
    st.session_state.quality_issues = quality_issues
    
    # Calculate overall metrics
    total_sales = st.session_state.sales_data['quantity'].sum()
    total_returns = st.session_state.returns_data['quantity'].sum()
    overall_return_rate = (total_returns / total_sales * 100) if total_sales > 0 else 0
    
    st.session_state.metrics = {
        'overall_return_rate': overall_return_rate,
        'total_sales': total_sales,
        'total_returns': total_returns,
        'products_affected': return_summary['sku'].nunique(),
        'quality_issues': len(quality_issues),
        'period': 'All Time',
        'return_summary': return_summary,
        'velocity': velocity
    }

def display_return_trend(returns_df):
    """Display return rate trend chart"""
    
    # Group by month
    returns_df['month'] = pd.to_datetime(returns_df['return_date']).dt.to_period('M')
    monthly_returns = returns_df.groupby('month')['quantity'].sum().reset_index()
    monthly_returns['month'] = monthly_returns['month'].astype(str)
    
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=monthly_returns['month'],
        y=monthly_returns['quantity'],
        mode='lines+markers',
        name='Returns',
        line=dict(color='#dc2626', width=3),
        marker=dict(size=8)
    ))
    
    fig.update_layout(
        title="Monthly Return Volume",
        xaxis_title="Month",
        yaxis_title="Units Returned",
        height=400,
        hovermode='x unified'
    )
    
    st.plotly_chart(fig, use_container_width=True)

def display_detailed_analysis():
    """Display detailed analysis tab"""
    
    st.markdown('<div class="section-header"><h2>📊 Return Rate by Product</h2></div>', unsafe_allow_html=True)
    
    if 'return_summary' in st.session_state.metrics:
        summary = st.session_state.metrics['return_summary']
        
        # Sort by return rate
        summary = summary.sort_values('return_rate', ascending=False)
        
        # Create bar chart
        fig = px.bar(
            summary.head(20),
            x='sku',
            y='return_rate',
            title='Top 20 Products by Return Rate',
            labels={'return_rate': 'Return Rate (%)', 'sku': 'SKU'},
            color='return_rate',
            color_continuous_scale=['green', 'yellow', 'red']
        )
        
        fig.update_layout(height=500)
        st.plotly_chart(fig, use_container_width=True)
        
        # Display detailed table
        st.subheader("Detailed Product Metrics")
        
        # Add velocity data if available
        if 'velocity' in st.session_state.metrics:
            summary = pd.merge(summary, st.session_state.metrics['velocity'], on='sku', how='left')
        
        # Format for display
        display_cols = ['sku', 'total_sold', 'total_returned', 'return_rate', 'daily_avg_30d', 'daily_avg_90d']
        available_cols = [col for col in display_cols if col in summary.columns]
        
        st.dataframe(
            summary[available_cols].style.format({
                'return_rate': '{:.1f}%',
                'daily_avg_30d': '{:.1f}',
                'daily_avg_90d': '{:.1f}'
            }).background_gradient(subset=['return_rate'], cmap='RdYlGn_r'),
            use_container_width=True
        )

def display_capa_form():
    """Display CAPA form"""
    
    st.markdown('<div class="section-header"><h2>📝 CAPA Information</h2></div>', unsafe_allow_html=True)
    
    with st.form("capa_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            capa_number = st.text_input("CAPA Number", value=f"CAPA-{datetime.now().strftime('%Y%m%d')}-001")
            product = st.text_input("Product Name", placeholder="e.g., Wheelchair Bag Advanced")
            sku = st.text_input("Primary SKU", placeholder="e.g., LVA3100BLK")
            prepared_by = st.text_input("Prepared By", value="Quality Team")
        
        with col2:
            date = st.date_input("Date", value=datetime.now())
            department = st.selectbox("Department", ["Quality", "Engineering", "Production", "Supply Chain"])
            severity = st.selectbox("Severity", ["Critical", "Major", "Minor"])
            priority = st.selectbox("Priority", ["High", "Medium", "Low"])
        
        st.markdown("### Issue Description")
        issue_description = st.text_area(
            "Describe the issue",
            height=150,
            placeholder="Provide detailed description of the quality issue, including impact on customers and products"
        )
        
        st.markdown("### Root Cause Analysis")
        root_cause = st.text_area(
            "Root Cause Analysis",
            height=150,
            placeholder="Describe investigation findings and identified root causes"
        )
        
        st.markdown("### Corrective Actions")
        corrective_action = st.text_area(
            "Proposed Corrective Actions",
            height=150,
            placeholder="Detail immediate and long-term corrective actions"
        )
        
        st.markdown("### Preventive Actions")
        preventive_action = st.text_area(
            "Proposed Preventive Actions",
            height=150,
            placeholder="Describe actions to prevent recurrence"
        )
        
        submitted = st.form_submit_button("💾 Save CAPA Data", type="primary")
        
        if submitted:
            st.session_state.capa_data = {
                'capa_number': capa_number,
                'date': date.strftime('%Y-%m-%d'),
                'product': product,
                'sku': sku,
                'prepared_by': prepared_by,
                'department': department,
                'severity': severity,
                'priority': priority,
                'issue_description': issue_description,
                'root_cause': root_cause,
                'corrective_action': corrective_action,
                'preventive_action': preventive_action
            }
            st.success("✅ CAPA data saved successfully!")

def display_document_generation():
    """Display document generation tab"""
    
    st.markdown('<div class="section-header"><h2>📄 Generate CAPA Document</h2></div>', unsafe_allow_html=True)
    
    if not st.session_state.capa_data:
        st.warning("⚠️ Please fill out the CAPA form first")
        return
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("Document Options")
        
        use_ai = st.checkbox("Use AI to enhance document", value=True)
        if use_ai:
            ai_provider = st.radio("AI Provider", ["Anthropic (Claude)", "OpenAI (GPT-4)"])
        
        include_metrics = st.checkbox("Include metrics summary", value=True)
        include_charts = st.checkbox("Include charts", value=True)
        include_screenshots = st.checkbox("Include uploaded screenshots", value=True)
    
    with col2:
        st.subheader("Export Format")
        export_format = st.radio("Format", ["Word Document (.docx)", "PDF (coming soon)"])
    
    if st.button("🚀 Generate Document", type="primary"):
        generate_capa_document(use_ai, ai_provider if use_ai else None, include_metrics, include_charts, include_screenshots)

def generate_capa_document(use_ai, ai_provider, include_metrics, include_charts, include_screenshots):
    """Generate the CAPA document"""
    
    with st.spinner("Generating document..."):
        generator = AIDocumentGenerator()
        
        # Generate content
        if use_ai:
            quality_issues = st.session_state.quality_issues if 'quality_issues' in st.session_state else pd.DataFrame()
            screenshot_data = st.session_state.get('screenshot_data', None)
            
            content = generator.generate_capa_document(
                st.session_state.capa_data,
                st.session_state.metrics,
                quality_issues,
                screenshot_data,
                use_anthropic=(ai_provider == "Anthropic (Claude)")
            )
            
            if content:
                st.success("✅ AI content generated successfully!")
            else:
                st.warning("⚠️ AI generation failed, using template format")
                content = generate_template_content()
        else:
            content = generate_template_content()
        
        # Export to Word
        doc_buffer = generator.export_to_docx(
            content, 
            st.session_state.capa_data,
            include_screenshots=include_screenshots,
            screenshots=st.session_state.screenshots if include_screenshots else []
        )
        
        if doc_buffer:
            st.download_button(
                label="📥 Download CAPA Document",
                data=doc_buffer,
                file_name=f"CAPA_{st.session_state.capa_data['capa_number']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Display preview
            st.markdown("### Document Preview")
            st.text_area("Content", content, height=500)

def generate_template_content():
    """Generate basic template content without AI"""
    
    capa = st.session_state.capa_data
    metrics = st.session_state.metrics
    
    content = f"""CORRECTIVE AND PREVENTIVE ACTION (CAPA) REPORT

1. ISSUE IDENTIFICATION

CAPA Number: {capa['capa_number']}
Date: {capa['date']}
Product: {capa['product']}
SKU: {capa['sku']}
Severity: {capa['severity']}
Priority: {capa['priority']}

Issue Description:
{capa['issue_description']}

2. METRICS SUMMARY

Overall Return Rate: {metrics.get('overall_return_rate', 0):.2f}%
Total Units Returned: {metrics.get('total_returns', 0):,}
Total Units Sold: {metrics.get('total_sales', 0):,}
Products Affected: {metrics.get('products_affected', 0)}

3. ROOT CAUSE ANALYSIS

{capa['root_cause']}

4. CORRECTIVE ACTIONS

{capa['corrective_action']}

5. PREVENTIVE ACTIONS

{capa['preventive_action']}

6. IMPLEMENTATION PLAN

[To be developed]

7. EFFECTIVENESS VERIFICATION

[To be determined]

Prepared by: {capa['prepared_by']}
Department: {capa['department']}
"""
    
    return content

# Initialize Google credentials on startup
if 'google_creds' not in st.session_state:
    st.session_state.google_creds = GoogleIntegration.authenticate()

if __name__ == "__main__":
    main()
